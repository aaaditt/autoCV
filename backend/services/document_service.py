"""
Document generation service.
Converts optimized resume text back to PDF and DOCX.
"""
import io
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_pdf(resume_text: str) -> bytes:
    """Convert resume text to a clean PDF."""
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch
    )

    styles = getSampleStyleSheet()

    # Custom styles
    normal_style = ParagraphStyle(
        'ResumeNormal',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        spaceAfter=4,
        fontName='Helvetica'
    )

    heading_style = ParagraphStyle(
        'ResumeHeading',
        parent=styles['Normal'],
        fontSize=12,
        leading=16,
        spaceBefore=10,
        spaceAfter=4,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#1D1D1F')
    )

    name_style = ParagraphStyle(
        'ResumeName',
        parent=styles['Normal'],
        fontSize=18,
        leading=22,
        spaceAfter=6,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#1D1D1F')
    )

    story = []
    lines = resume_text.split('\n')

    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 4))
            continue

        # Heuristic: first non-empty line is usually the name
        if i == 0 or (i <= 2 and len(line) < 50 and not any(c in line for c in ['@', '|', '•', '-'])):
            story.append(Paragraph(line, name_style))
        # Section headers: all caps or ends with colon
        elif line.isupper() or (line.endswith(':') and len(line) < 40):
            story.append(Paragraph(line, heading_style))
        else:
            # Replace bullet characters
            if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                line = '• ' + line.lstrip('•-* ')
            story.append(Paragraph(line, normal_style))

    doc.build(story)
    return buffer.getvalue()


def generate_docx(resume_text: str) -> bytes:
    """Convert resume text to a clean DOCX."""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    lines = resume_text.split('\n')

    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        # Name (first line heuristic)
        if i == 0 or (i <= 2 and len(line) < 50):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(18)
        # Section headers
        elif line.isupper() or (line.endswith(':') and len(line) < 40):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
        # Bullet points
        elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(line.lstrip('•-* '))
        else:
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
